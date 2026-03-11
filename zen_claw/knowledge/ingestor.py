"""Document ingestion for local files and URLs."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class Document:
    content: str
    source: str
    page: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


class Ingestor:
    SUPPORTED_TEXT_SUFFIXES = {".txt", ".md", ".rst"}
    SUPPORTED_HTML_SUFFIXES = {".html", ".htm"}
    SUPPORTED_SUFFIXES = {".pdf", ".docx", *SUPPORTED_TEXT_SUFFIXES, *SUPPORTED_HTML_SUFFIXES}

    def __init__(
        self,
        max_url_chars: int = 200_000,
        user_agent: str = "zen-claw-rag/0.1",
        http_timeout: float = 20.0,
    ):
        self.max_url_chars = max_url_chars
        self.user_agent = user_agent
        self.http_timeout = http_timeout

    async def ingest(self, source: str) -> list[Document]:
        if source.startswith(("http://", "https://")):
            return await self._ingest_url(source)
        path = Path(source)
        if path.is_dir():
            return self._ingest_directory(path)
        return self._ingest_path(path)

    def _ingest_path(self, path: Path) -> list[Document]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._ingest_pdf(path)
        if suffix == ".docx":
            return self._ingest_docx(path)
        if suffix in self.SUPPORTED_TEXT_SUFFIXES:
            return self._ingest_text(path)
        if suffix in self.SUPPORTED_HTML_SUFFIXES:
            return self._ingest_html(path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _ingest_directory(self, path: Path) -> list[Document]:
        docs: list[Document] = []
        for child in sorted(x for x in path.rglob("*") if x.is_file()):
            if child.suffix.lower() not in self.SUPPORTED_SUFFIXES:
                continue
            docs.extend(self._ingest_path(child))
        if docs:
            return docs
        raise ValueError(f"No supported files found in directory: {path}")

    def _ingest_pdf(self, path: Path) -> list[Document]:
        try:
            import fitz
        except ImportError as exc:
            raise ImportError("pymupdf is required for PDF ingestion") from exc
        docs: list[Document] = []
        pdf = fitz.open(str(path))
        try:
            for i in range(len(pdf)):
                text = (pdf[i].get_text("text") or "").strip()
                if text:
                    docs.append(
                        Document(
                            content=text,
                            source=str(path),
                            page=i + 1,
                            metadata={"total_pages": len(pdf)},
                        )
                    )
        finally:
            pdf.close()
        return docs

    def _ingest_docx(self, path: Path) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("python-docx is required for DOCX ingestion") from exc
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if not paragraphs:
            return []
        docs: list[Document] = []
        cur: list[str] = []
        cur_len = 0
        section = 1
        for para in paragraphs:
            if cur and cur_len + len(para) > 1000:
                docs.append(
                    Document(
                        content="\n\n".join(cur),
                        source=str(path),
                        page=section,
                        metadata={"section": section},
                    )
                )
                section += 1
                cur = [para]
                cur_len = len(para)
            else:
                cur.append(para)
                cur_len += len(para)
        if cur:
            docs.append(
                Document(
                    content="\n\n".join(cur),
                    source=str(path),
                    page=section,
                    metadata={"section": section},
                )
            )
        return docs

    def _ingest_text(self, path: Path) -> list[Document]:
        content = path.read_text(encoding="utf-8", errors="replace").strip()
        return [Document(content=content, source=str(path), page=None)]

    def _ingest_html(self, path: Path) -> list[Document]:
        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            import trafilatura

            extracted = trafilatura.extract(raw, include_comments=False) or raw
        except ImportError:
            extracted = raw
        return [Document(content=extracted.strip(), source=str(path), page=None)]

    async def _ingest_url(self, url: str) -> list[Document]:
        import httpx

        async with httpx.AsyncClient(
            timeout=self.http_timeout,
            follow_redirects=True,
            headers={"User-Agent": self.user_agent},
        ) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            body = resp.text[: self.max_url_chars]
        try:
            import trafilatura

            extracted = trafilatura.extract(body, include_comments=False) or body
        except ImportError:
            extracted = body
        return [Document(content=extracted.strip(), source=url, page=None)]
